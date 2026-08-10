import io

import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter

from app.rag.errors import ProcessingError
from app.rag.extractors import (
    DocxExtractor,
    ExtractionService,
    PdfExtractor,
    PlainTextExtractor,
)


def _make_pdf(*texts: str) -> bytes:
    writer = PdfWriter()
    for text in texts:
        writer.add_blank_page(width=400, height=600)
    buffer = io.BytesIO()
    # pypdf has no built-in text printer; build minimal valid PDF with text streams
    writer.write(buffer)
    return buffer.getvalue()


def test_plain_text_utf8_and_latin1():
    assert PlainTextExtractor().extract("héllo wörld".encode("utf-8")) == "héllo wörld"
    fallback = PlainTextExtractor().extract(bytes([0xE9, 0x20]) + b"x")
    assert "x" in fallback


def test_extraction_service_dispatches_plain_text():
    service = ExtractionService()
    assert service.extract("plain text".encode("utf-8"), "text/plain") == "plain text"


def test_extraction_service_rejects_unknown_type():
    service = ExtractionService()
    with pytest.raises(ProcessingError):
        service.extract(b"x", "application/octet-stream")


def test_extraction_service_raises_on_empty_text():
    service = ExtractionService()
    with pytest.raises(ProcessingError):
        service.extract(b"   \n\n ", "text/plain")


def test_docx_extractor():
    buffer = io.BytesIO()
    doc = DocxDocument()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    doc.save(buffer)
    text = DocxExtractor().extract(buffer.getvalue())
    assert "First paragraph" in text
    assert "Second paragraph" in text


def test_docx_extractor_rejects_corrupt():
    with pytest.raises(ProcessingError):
        DocxExtractor().extract(b"this is not a docx")


def test_pdf_extractor_rejects_corrupt():
    with pytest.raises(ProcessingError):
        PdfExtractor().extract(b"not a real pdf %PDF-1.0 garbage")